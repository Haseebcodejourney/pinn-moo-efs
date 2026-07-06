"""Build updated paper DOCX + PDF from LMNotes_full_paper.md with embedded figures."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
PAPER_MD = Path(__file__).resolve().parent / "LMNotes_full_paper.md"
FIG_DIR = ROOT / "results" / "figures"
OUT_DOCX = Path(__file__).resolve().parent / "PINN_MOO_EFS_paper_revised.docx"
OUT_PDF = Path(__file__).resolve().parent / "PINN_MOO_EFS_paper_revised.pdf"

# Figure paths referenced in markdown -> actual files
FIGURE_MAP = {
    "fig_workflow.png": FIG_DIR / "fig_workflow.png",
    "poisson_fig_pareto_f1_f2.png": FIG_DIR / "poisson_fig_pareto_f1_f2.png",
    "poisson_fig_pareto_f1_f3.png": FIG_DIR / "poisson_fig_pareto_f1_f3.png",
    "poisson_fig_solution.png": FIG_DIR / "poisson_fig_solution.png",
    "burgers_fig_pareto_f1_f2.png": FIG_DIR / "burgers_fig_pareto_f1_f2.png",
    "burgers_fig_pareto_f1_f3.png": FIG_DIR / "burgers_fig_pareto_f1_f3.png",
    "burgers_fig_solution.png": FIG_DIR / "burgers_fig_solution.png",
    "burgers_fig_efs_comparison.png": FIG_DIR / "burgers_fig_efs_comparison.png",
    "fig_data_scarcity.png": FIG_DIR / "fig_data_scarcity.png",
}


def strip_md_bold(text: str) -> tuple[str, bool]:
    text = text.strip()
    if text.startswith("**") and text.endswith("**"):
        return text[2:-2], True
    return text, False


def add_rich_paragraph(doc: Document, text: str, style: str | None = None, bold: bool = False):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
            if bold:
                run.bold = True
    return p


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    header = [c.strip() for c in lines[0].strip("|").split("|")]
    rows = []
    for line in lines[2:]:
        if not line.strip().startswith("|"):
            break
        rows.append([c.strip() for c in line.strip("|").split("|")])
    return header, rows


def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val.replace("**", "")
    doc.add_paragraph()


def embed_figures_from_line(doc: Document, line: str):
    for name, path in FIGURE_MAP.items():
        if name in line and path.exists():
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run()
            run.add_picture(str(path), width=Inches(5.8))
            doc.add_paragraph(name, style="Caption")
            doc.add_paragraph()


def build_docx() -> Path:
    text = PAPER_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            t = line[2:].strip()
            h = doc.add_heading(t, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            block = [line]
            j = i + 1
            while j < len(lines) and lines[j].startswith("|"):
                block.append(lines[j].rstrip())
                j += 1
            header, rows = parse_table(block)
            add_table(doc, header, rows)
            i = j
            continue

        if line.startswith("**Figure") or line.startswith("*Insert") or "results/figures/" in line:
            add_rich_paragraph(doc, line.lstrip("*").strip())
            embed_figures_from_line(doc, line)
            i += 1
            continue

        if line.startswith("- ") or line.startswith("1. "):
            add_rich_paragraph(doc, line[2:].strip() if line.startswith("- ") else line, style="List Bullet")
            i += 1
            continue

        if line.startswith("**") and line.endswith("**") and ":" in line:
            key, _, val = line.partition(":")
            p = doc.add_paragraph()
            r1 = p.add_run(key.replace("**", "") + ":")
            r1.bold = True
            p.add_run(val)
            i += 1
            continue

        if line.startswith("$$"):
            eq = line[2:-2] if line.endswith("$$") else line[2:]
            p = doc.add_paragraph(eq.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("*Reproduce:"):
            add_rich_paragraph(doc, line.strip("*"))
            i += 1
            continue

        add_rich_paragraph(doc, line)
        i += 1

    doc.save(OUT_DOCX)
    return OUT_DOCX


def export_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        # 17 = wdExportFormatPDF
        doc.ExportAsFixedFormat(str(pdf_path.resolve()), ExportFormat=17)
        doc.Close(False)
        word.Quit()
        return pdf_path.exists()
    except Exception as exc:
        print(f"Word PDF export failed: {exc}", file=sys.stderr)
        return False


def main():
    if not PAPER_MD.exists():
        raise FileNotFoundError(PAPER_MD)
    missing = [n for n, p in FIGURE_MAP.items() if not p.exists()]
    if missing:
        print("Warning: missing figures:", ", ".join(missing))

    docx = build_docx()
    print(f"DOCX -> {docx}")

    if export_pdf(docx, OUT_PDF):
        print(f"PDF  -> {OUT_PDF}")
    else:
        print("Open the DOCX in Word and use Save As -> PDF if auto-export failed.")


if __name__ == "__main__":
    main()
